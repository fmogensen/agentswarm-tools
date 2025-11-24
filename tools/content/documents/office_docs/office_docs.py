"""
Generate professional Word documents (.docx) from structured content
"""

import os
import re
import tempfile
import urllib.parse
import urllib.request
from datetime import datetime
from typing import Any, ClassVar, Dict, List, Optional

from pydantic import Field

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None

from shared.base import BaseTool
from shared.errors import APIError, ConfigurationError, ValidationError


class OfficeDocsTool(BaseTool):
    """
    Generate or modify professional Word documents (.docx) from structured content.

    Use this tool to create or modify formatted documents including reports, proposals,
    memos, letters, and other business documents with proper styling.

    Args:
        mode: Operation mode - "create" to create new document, "modify" to update existing
        content: Document content as text or structured data (supports markdown)
        template: Template type (report, proposal, memo, letter, blank) - only for create mode
        title: Document title
        include_toc: Whether to include table of contents
        font_name: Font family (Calibri, Arial, Times New Roman)
        font_size: Base font size in points (default: 11)
        output_format: Output format (docx, pdf, both)
        existing_file_url: URL to existing document (required for modify mode)

    Returns:
        Dict containing:
        - success: Boolean indicating success
        - result: {"document_url": "...", "format": "...", "pages": N, "mode": "..."}
        - metadata: Tool execution metadata

    Example (Create):
        >>> tool = OfficeDocsTool(
        ...     mode="create",
        ...     content="# Report\\n\\nThis is the content...",
        ...     template="report",
        ...     title="Q4 Sales Report",
        ...     include_toc=True
        ... )
        >>> result = tool.run()
        >>> print(result['result']['document_url'])

    Example (Modify):
        >>> tool = OfficeDocsTool(
        ...     mode="modify",
        ...     existing_file_url="computer:///path/to/doc.docx",
        ...     content="\\n\\n# New Section\\n\\nAdditional content...",
        ...     title="Updated Q4 Sales Report"
        ... )
        >>> result = tool.run()
    """

    # Tool metadata
    tool_name: str = "office_docs"
    tool_category: str = "content"

    # Parameters
    mode: str = Field("create", description="Operation mode: create or modify")
    content: str = Field(..., description="Document content (supports markdown)", min_length=1)
    template: str = Field(
        "blank",
        description="Template type: report, proposal, memo, letter, blank (create mode only)",
    )
    title: str = Field("Untitled Document", description="Document title")
    include_toc: bool = Field(False, description="Include table of contents")
    font_name: str = Field("Calibri", description="Font family")
    font_size: int = Field(11, description="Base font size in points", ge=8, le=24)
    output_format: str = Field("docx", description="Output format: docx, pdf, both")
    existing_file_url: Optional[str] = Field(
        None, description="URL to existing file (required for modify mode)"
    )

    # Allowed values
    ALLOWED_MODES: ClassVar[List[str]] = ["create", "modify"]
    ALLOWED_TEMPLATES: ClassVar[List[str]] = ["report", "proposal", "memo", "letter", "blank"]
    ALLOWED_FORMATS: ClassVar[List[str]] = ["docx", "pdf", "both"]
    ALLOWED_FONTS: ClassVar[List[str]] = [
        "Calibri",
        "Arial",
        "Times New Roman",
        "Georgia",
        "Verdana",
    ]

    def _execute(self) -> Dict[str, Any]:
        """
        Execute the office_docs tool.

        Returns:
            Dict with results
        """
        # 1. CHECK MOCK MODE FIRST (before validation that requires libraries)

        self._logger.info(
            f"Executing {self.tool_name} with mode={self.mode}, content={self.content}, template={self.template}, ..."
        )
        if self._should_use_mock():
            # Still do basic parameter validation
            self._validate_basic_parameters()
            return self._generate_mock_results()

        # 2. FULL VALIDATION (including library checks)
        self._validate_parameters()

        # 3. EXECUTE
        try:
            result = self._process()

            self._logger.info(f"Successfully completed {self.tool_name}")

            return {
                "success": True,
                "result": result,
                "metadata": {"tool_name": self.tool_name},
            }
        except Exception as e:
            self._logger.error(f"Error in {self.tool_name}: {str(e)}", exc_info=True)
            raise APIError(f"Failed: {e}", tool_name=self.tool_name)

    def _validate_basic_parameters(self) -> None:
        """Validate basic parameters (used in mock mode)."""
        # Validate mode
        if self.mode not in self.ALLOWED_MODES:
            raise ValidationError(
                f"Invalid mode '{self.mode}'. Must be one of: {', '.join(self.ALLOWED_MODES)}",
                tool_name=self.tool_name,
                field="mode",
            )

        # Validate mode-specific requirements
        if self.mode == "modify":
            if not self.existing_file_url:
                raise ValidationError(
                    "existing_file_url is required when mode='modify'",
                    tool_name=self.tool_name,
                    field="existing_file_url",
                )
        elif self.mode == "create":
            if self.existing_file_url:
                raise ValidationError(
                    "existing_file_url should not be provided when mode='create'",
                    tool_name=self.tool_name,
                    field="existing_file_url",
                )

        # Validate content
        if not self.content.strip():
            raise ValidationError("Content cannot be empty", tool_name=self.tool_name)

        # Validate template (only matters for create mode)
        if self.mode == "create" and self.template not in self.ALLOWED_TEMPLATES:
            raise ValidationError(
                f"Invalid template '{self.template}'. Must be one of: {', '.join(self.ALLOWED_TEMPLATES)}",
                tool_name=self.tool_name,
                field="template",
            )

        # Validate output format
        if self.output_format not in self.ALLOWED_FORMATS:
            raise ValidationError(
                f"Invalid output_format '{self.output_format}'. Must be one of: {', '.join(self.ALLOWED_FORMATS)}",
                tool_name=self.tool_name,
                field="output_format",
            )

        # Validate font
        if self.font_name not in self.ALLOWED_FONTS:
            raise ValidationError(
                f"Invalid font_name '{self.font_name}'. Must be one of: {', '.join(self.ALLOWED_FONTS)}",
                tool_name=self.tool_name,
                field="font_name",
            )

    def _validate_parameters(self) -> None:
        """Validate input parameters (full validation including dependencies)."""
        # Do basic validation first
        self._validate_basic_parameters()

        # Check if python-docx is installed
        if Document is None:
            raise ConfigurationError(
                "python-docx library not installed. Run: pip install python-docx",
                tool_name=self.tool_name,
            )

    def _should_use_mock(self) -> bool:
        """Check if mock mode enabled."""
        return os.getenv("USE_MOCK_APIS", "false").lower() == "true"

    def _generate_mock_results(self) -> Dict[str, Any]:
        """Generate mock results for testing."""
        return {
            "success": True,
            "result": {
                "document_url": "https://mock.example.com/doc123.docx",
                "format": self.output_format,
                "pages": 5,
                "file_size": "245 KB",
                "title": self.title,
                "template": self.template if self.mode == "create" else "N/A",
                "mode": self.mode,
            },
            "metadata": {"mock_mode": True},
        }

    def _process(self) -> Dict[str, Any]:
        """Main processing logic."""
        if self.mode == "create":
            return self._process_create()
        else:  # modify
            return self._process_modify()

    def _process_create(self) -> Dict[str, Any]:
        """Create new document from scratch."""
        # Create Document object
        doc = Document()

        # Apply template-specific styling
        self._apply_template_styles(doc)

        # Add title
        self._add_title(doc)

        # Add table of contents if requested
        if self.include_toc:
            self._add_toc_placeholder(doc)

        # Parse and add content
        self._parse_and_add_content(doc)

        # Save and return result
        return self._save_and_return(doc)

    def _process_modify(self) -> Dict[str, Any]:
        """Modify existing document."""
        # Download existing file
        local_path = self._download_file(self.existing_file_url)

        try:
            # Open existing document
            doc = Document(local_path)

            # Update title if different from default
            if self.title != "Untitled Document":
                # Find and update first heading or add new title
                title_updated = False
                for para in doc.paragraphs:
                    if para.style.name == "Heading 1":
                        para.text = self.title
                        title_updated = True
                        break

                if not title_updated:
                    # Add title at beginning
                    new_para = doc.paragraphs[0].insert_paragraph_before(
                        self.title, style="Heading 1"
                    )

            # Apply font changes to existing content if specified
            if self.font_name != "Calibri" or self.font_size != 11:
                self._update_document_fonts(doc)

            # Add table of contents if requested and not present
            if self.include_toc:
                # Check if TOC already exists
                has_toc = any("Table of Contents" in p.text for p in doc.paragraphs)
                if not has_toc:
                    self._add_toc_placeholder(doc)

            # Append new content
            self._parse_and_add_content(doc)

            # Save and return result
            return self._save_and_return(doc)

        finally:
            # Clean up downloaded file
            try:
                os.unlink(local_path)
            except:
                pass

    def _download_file(self, url: str) -> str:
        """Download file from URL to temporary location."""
        # Handle computer:// protocol
        if url.startswith("computer://"):
            file_path = url.replace("computer://", "")
            if os.path.exists(file_path):
                return file_path
            else:
                raise ValidationError(f"File not found: {file_path}", tool_name=self.tool_name)

        # Handle http/https URLs
        elif url.startswith("http://") or url.startswith("https://"):
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            temp_file.close()

            try:
                urllib.request.urlretrieve(url, temp_file.name)
                return temp_file.name
            except Exception as e:
                self._logger.error(f"Error in {self.tool_name}: {str(e)}", exc_info=True)
                raise APIError(f"Failed to download file from URL: {e}", tool_name=self.tool_name)

        else:
            raise ValidationError(f"Unsupported URL scheme: {url}", tool_name=self.tool_name)

    def _update_document_fonts(self, doc: Any) -> None:
        """Update fonts throughout existing document."""
        # Update Normal style
        styles = doc.styles
        if "Normal" in styles:
            normal_style = styles["Normal"]
            normal_font = normal_style.font
            normal_font.name = self.font_name
            normal_font.size = Pt(self.font_size)

        # Update all paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = self.font_name
                # Only update size if not a heading
                if para.style.name not in ["Heading 1", "Heading 2", "Heading 3"]:
                    run.font.size = Pt(self.font_size)

    def _save_and_return(self, doc: Any) -> Dict[str, Any]:
        """Save document and return result metadata."""
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()

        # Generate accessible URL (in production, upload to storage)
        document_url = self._upload_to_storage(temp_file.name)

        # Count pages (approximate based on paragraphs)
        estimated_pages = max(1, len(doc.paragraphs) // 25)

        # Get file size
        file_size_bytes = os.path.getsize(temp_file.name)
        file_size = self._format_file_size(file_size_bytes)

        # Clean up temp file
        try:
            os.unlink(temp_file.name)
        except:
            pass

        return {
            "document_url": document_url,
            "format": self.output_format,
            "pages": estimated_pages,
            "file_size": file_size,
            "title": self.title,
            "template": self.template if self.mode == "create" else "N/A",
            "mode": self.mode,
        }

    def _apply_template_styles(self, doc: Any) -> None:
        """Apply template-specific styles to the document."""
        # Get or create styles
        styles = doc.styles

        # Configure Normal style
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = self.font_name
        normal_font.size = Pt(self.font_size)

        # Configure Heading 1
        if "Heading 1" in styles:
            h1_style = styles["Heading 1"]
            h1_font = h1_style.font
            h1_font.name = self.font_name
            h1_font.size = Pt(self.font_size + 6)
            h1_font.bold = True

        # Configure Heading 2
        if "Heading 2" in styles:
            h2_style = styles["Heading 2"]
            h2_font = h2_style.font
            h2_font.name = self.font_name
            h2_font.size = Pt(self.font_size + 4)
            h2_font.bold = True

        # Template-specific styling
        if self.template == "memo":
            # Memo typically uses tighter spacing
            normal_style.paragraph_format.space_after = Pt(6)
        elif self.template == "letter":
            # Letter uses standard spacing
            normal_style.paragraph_format.space_after = Pt(8)
        elif self.template in ["report", "proposal"]:
            # Reports and proposals use more spacing
            normal_style.paragraph_format.space_after = Pt(12)

    def _add_title(self, doc: Any) -> None:
        """Add title to the document."""
        if self.template == "memo":
            # Memo header format
            doc.add_paragraph("MEMORANDUM", style="Heading 1")
            doc.add_paragraph(f"Subject: {self.title}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph()
        elif self.template == "letter":
            # Letter header format
            doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            doc.add_paragraph()
            doc.add_paragraph(self.title, style="Heading 1")
            doc.add_paragraph()
        else:
            # Standard title
            title_para = doc.add_paragraph(self.title, style="Heading 1")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    def _add_toc_placeholder(self, doc: Any) -> None:
        """Add a placeholder for table of contents."""
        # Add TOC header
        toc_para = doc.add_paragraph("Table of Contents", style="Heading 2")
        doc.add_paragraph()

        # Add note about TOC
        note = doc.add_paragraph()
        note.add_run("[Table of Contents will be generated when document is opened in Word. ")
        note.add_run("Right-click and select 'Update Field' to refresh.]")
        note.runs[0].italic = True
        note.runs[1].italic = True

        doc.add_paragraph()

    def _parse_and_add_content(self, doc: Any) -> None:
        """Parse content (supports basic markdown) and add to document."""
        lines = self.content.split("\n")

        i = 0
        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue

            # Heading 1 (# Heading)
            if line.startswith("# "):
                heading_text = line[2:].strip()
                doc.add_paragraph(heading_text, style="Heading 1")
                i += 1
                continue

            # Heading 2 (## Heading)
            if line.startswith("## "):
                heading_text = line[3:].strip()
                doc.add_paragraph(heading_text, style="Heading 2")
                i += 1
                continue

            # Heading 3 (### Heading)
            if line.startswith("### "):
                heading_text = line[4:].strip()
                para = doc.add_paragraph(heading_text)
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(self.font_size + 2)
                i += 1
                continue

            # Bullet point (- item or * item)
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                bullet_text = line.strip()[2:].strip()
                doc.add_paragraph(bullet_text, style="List Bullet")
                i += 1
                continue

            # Numbered list (1. item)
            if re.match(r"^\d+\.\s", line.strip()):
                list_text = re.sub(r"^\d+\.\s", "", line.strip())
                doc.add_paragraph(list_text, style="List Number")
                i += 1
                continue

            # Table detection (basic pipe tables)
            if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
                # Try to parse as table
                table_lines = [line]
                j = i + 1
                while j < len(lines) and "|" in lines[j]:
                    table_lines.append(lines[j])
                    j += 1

                # Add table to document
                self._add_table(doc, table_lines)
                i = j
                continue

            # Regular paragraph
            doc.add_paragraph(line)
            i += 1

    def _add_table(self, doc: Any, table_lines: list) -> None:
        """Add a table to the document from markdown-style table lines."""
        # Parse table data
        rows = []
        for line in table_lines:
            # Skip separator lines (|---|---|)
            if re.match(r"^[\s\|:\-]+$", line):
                continue

            # Split by pipe and clean
            cells = [cell.strip() for cell in line.split("|")]
            # Remove empty first/last cells if present
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]

            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Light Grid Accent 1"

        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text

                    # Make header row bold
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _upload_to_storage(self, file_path: str) -> str:
        """
        Upload file to storage and return accessible URL.

        In production, this would upload to AI Drive or cloud storage.
        For now, returns a file:// URL for local testing.
        """
        # Check for storage configuration
        aidrive_key = os.getenv("AIDRIVE_API_KEY")

        if aidrive_key:
            # In production: upload to AI Drive
            # For now, return mock URL
            filename = os.path.basename(file_path)
            return f"https://aidrive.example.com/documents/{filename}"
        else:
            # Return local file path as computer:// URL
            return f"computer://{file_path}"

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format."""
        for unit in ["B", "KB", "MB", "GB"]:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"


if __name__ == "__main__":
    # Test the office_docs tool
    print("Testing OfficeDocsTool...")

    # Test with mock mode
    import os

    os.environ["USE_MOCK_APIS"] = "true"

    # Test basic document
    tool = OfficeDocsTool(
        content="# Test Report\n\nThis is a test document.",
        template="report",
        title="Test Document",
    )
    result = tool.run()

    assert result.get("success") == True
    assert "document_url" in result["result"]
    print(f"✅ Basic test passed")

    # Test with table of contents
    tool2 = OfficeDocsTool(
        content="# Chapter 1\n\nContent...\n\n# Chapter 2\n\nMore content...",
        include_toc=True,
        output_format="pdf",
    )
    result2 = tool2.run()
    assert result2.get("success") == True
    print(f"✅ TOC test passed")

    # Test memo template
    tool3 = OfficeDocsTool(
        content="This is important information.", template="memo", title="Important Update"
    )
    result3 = tool3.run()
    assert result3.get("success") == True
    print(f"✅ Memo template test passed")

    # Test with markdown formatting
    tool4 = OfficeDocsTool(
        content="""# Introduction

This is a report with **formatting**.

## Section 1

- Bullet point 1
- Bullet point 2
- Bullet point 3

## Section 2

1. First item
2. Second item
3. Third item

Regular paragraph text here.""",
        template="report",
        title="Formatted Report",
        font_name="Arial",
        font_size=12,
    )
    result4 = tool4.run()
    assert result4.get("success") == True
    print(f"✅ Markdown formatting test passed")

    # Test modify mode validation
    tool5 = OfficeDocsTool(
        mode="modify",
        content="New content",
        # Missing existing_file_url - should fail
    )
    result5 = tool5.run()
    assert result5.get("success") == False  # Should fail validation
    assert "existing_file_url is required" in str(result5.get("error", ""))
    print(f"✅ Modify mode validation test passed")

    # Test create mode with existing_file_url (should fail)
    tool6 = OfficeDocsTool(
        mode="create",
        content="New content",
        existing_file_url="computer:///some/file.docx",  # Should not be provided for create
    )
    result6 = tool6.run()
    assert result6.get("success") == False  # Should fail validation
    assert "should not be provided" in str(result6.get("error", ""))
    print(f"✅ Create mode validation test passed")

    # Test modify mode with mock
    tool7 = OfficeDocsTool(
        mode="modify",
        content="# New Section\n\nAdditional content",
        existing_file_url="https://mock.example.com/doc.docx",
        title="Updated Document",
    )
    result7 = tool7.run()
    assert result7.get("success") == True
    assert result7["result"]["mode"] == "modify"
    print(f"✅ Modify mode mock test passed")

    print("\n🎉 All tests passed!")
